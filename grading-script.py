import os
import json
import time
import docx
import re
import pandas as pd
from pathlib import Path
from datetime import datetime
import PyPDF2
from openai import OpenAI
import statistics
from typing import List, Tuple, Optional
import argparse

# --- API Configuration ---
API_KEY = "your-openai-key"
BASE_URL = "https://api.openai.com/v1"

# --- Model Configuration ---
GRADING_MODEL = "gpt-4.1"
TEMPERATURE = 0.35
MAX_ATTEMPTS = 5  # Maximum number of retries for grading a single manuscript

# Initialize OpenAI client
client = OpenAI(
    api_key=API_KEY,
    base_url=BASE_URL
)

# Input and output directories
INPUT_FOLDER = "./input"
OUTPUT_FILE = "manuscript_grading_results.xlsx"
DEFAULT_OUTPUT_FILE = "manuscript_grading_results.xlsx"


def get_system_prompt() -> str:
    """Create the system prompt with grading instructions"""
    return """You are tasked with grading manuscripts in the intersection of Chemistry, Solid State Physics and Material Science. Six prestige levels of well-known journals are listed below for reference, but you do not need to consider manuscript fit to specific journals; these are to convey the relative prestige of each level. For each manuscript, provide your educated guess—expressed as a percentage—for the chance it would be sent out for external review at each of the six journal levels. In making your estimates, consider overall quality, scope breadth, methodological novelty, interest to readership, workload, quality of writing, methodological rigour, and whether the results fully support the claims. Some manuscripts you grade may already be published articles, but please evaluate them as if they are new, without regard to where they were actually published. Note each lower level should have a equal or higher chance than the previous level.
*Level 1 - Very top journal*
Nature; Science.
*Level 2 - Well-regarded sister journals of the very top journals*
Nature Materials; Nature Chemistry; Nature Electronic; Nature Physics; Nature Energy; Nature Nanotechnology.
*Level 3 - Weaker sister journals of the very top journals or equivalent*
Nature Communications; Science Advances; Joule.
*Level 4 - Top broad field journals*
Journal of the American Chemical Society; Advanced Materials; Angewandte Chemie International Edition; Physical Review Letters; Physical Review X; PNAS; Chem.
*Level 5 - Good broad field journals and equivalent*
JACS Au; Advanced Functional Materials; Advanced Energy Materials; CCS Chemistry; Chemical Science; ACS Central Science; Advanced Science; Energy and Environmental Science; Physical Review B
*Level 6 - Good narrow field journals*
Journal of Materials Chemistry A; Small; Chemical Communications; Macromolecules; Advanced Electronic Materials; Journal of Physical Chemistry Letters; Journal of Physical Chemistry A; ACS Applied Materials & Interface.
## Output Format
Your output must be a JSON object with:
- Keys Level 1 through Level 6, each mapping to an integer percentage (0–100) indicating your estimate of the chance the manuscript is sent for external review at that level.
- A single key "justification" with a one-sentence rationale for your scoring.
Example output:
{
  "Level 1": 0,
  "Level 2": 10,
  "Level 3": 50,
  "Level 4": 80,
  "Level 5": 90,
  "Level 6": 100,
  "justification": "The methodology is solid, but the novelty and breadth do not meet the expectations for the highest-impact journals."
}"""


def get_user_prompt(manuscript_text: str) -> str:
    """Create the user prompt with manuscript text"""
    return f"Manuscript to grade:\n\n{manuscript_text}"


def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from a DOCX file"""
    try:
        doc = docx.Document(file_path)
        full_text = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)

        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading {file_path}: {str(e)}")
        return ""

def parse_args():
    """Parse command-line arguments."""
    parser = argparse.ArgumentParser(
        description="Batch Manuscript Grader: specify an input folder of .docx/.txt/.pdf files."
    )
    parser.add_argument(
        'input_folder',
        nargs='?',
        default='./input',
        help="Path to folder containing manuscripts (default: ./input)"
    )
    parser.add_argument(
        '--output-file',
        '-o',
        default=DEFAULT_OUTPUT_FILE,
        help=f"Excel filename to save results into (default: {DEFAULT_OUTPUT_FILE})"
    )
    return parser.parse_args()


def extract_text_from_txt(file_path: str) -> str:
    """Extract all text from a TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading {file_path}: {str(e)}")
            return ""
    except Exception as e:
        print(f"Error reading {file_path}: {str(e)}")
        return ""


def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text from a PDF file"""
    try:
        text_content = []

        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            # Extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    text_content.append(text)

        extracted_text = "\n".join(text_content)

        # Basic cleanup - remove excessive whitespace
        extracted_text = re.sub(r'\n\s*\n', '\n\n', extracted_text)  # Multiple newlines to double newline
        extracted_text = re.sub(r' +', ' ', extracted_text)  # Multiple spaces to single space

        return extracted_text

    except Exception as e:
        print(f"Error reading PDF {file_path}: {str(e)}")
        print("Note: Make sure PyPDF2 is installed: pip install PyPDF2")
        return ""


def extract_text_from_file(file_path: str) -> str:
    """Extract text from DOCX, TXT, or PDF file based on extension"""
    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.txt':
        return extract_text_from_txt(file_path)
    elif file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    else:
        print(f"Unsupported file type: {file_extension}")
        return ""


def grade_manuscript(manuscript_text: str):
    """Send manuscript to grading model for grading and return parsed scores"""
    try:
        # Prepare messages with system and user prompts
        messages = [
            {"role": "system", "content": get_system_prompt()},
            {"role": "user", "content": get_user_prompt(manuscript_text)}
        ]

        response = client.chat.completions.create(
            model=GRADING_MODEL,
            messages=messages,
            temperature=TEMPERATURE
        )

        result = response.choices[0].message.content

        # Wait 1 second before next API call
        time.sleep(1)

        # Parse JSON response directly
        try:
            grading_data = json.loads(result)

            # Extract scores in order [Level 1, Level 2, ..., Level 6] and justification
            scores = []
            for i in range(1, 7):
                level_key = f"Level {i}"
                score = grading_data.get(level_key)
                # Ensure score is treated as a number, handle potential string values
                if score is not None:
                    scores.append(float(score))
                else:
                    scores.append(None)

            justification = grading_data.get("justification", "No justification provided")

            return scores, justification

        except (json.JSONDecodeError, ValueError, KeyError) as e:
            print(f"Error parsing JSON response: {str(e)}")
            print(f"Raw response: {result}")
            return [None] * 6, "Error parsing response"

    except Exception as e:
        print(f"Error in grading: {str(e)}")
        time.sleep(1)  # Wait even on error
        return [None] * 6, "Error in API call"


def calculate_weighted_score(scores):
    """
    Calculate weighted average: level 1 weight 4, level 2 weight 2, others weight 1

    Args:
        scores: List of 6 scores [level1, level2, level3, level4, level5, level6]

    Returns:
        float: Weighted average score
    """
    weights = [2, 2, 1, 1, 1, 1]
    weighted_sum = 0
    total_weight = 0

    for i, score in enumerate(scores):
        if score is not None:
            weighted_sum += score * weights[i]
            total_weight += weights[i]

    if total_weight == 0:
        return 0

    return weighted_sum / total_weight


def are_scores_ascending(scores: List[Optional[float]]) -> bool:
    """
    Check if scores are in ascending order.
    Each subsequent level should have an equal or higher chance than the previous one.
    """
    # If any scores are missing, we cannot validate, so treat as failure.
    if any(s is None for s in scores):
        return False

    for i in range(len(scores) - 1):
        if scores[i] > scores[i + 1]:
            return False
    return True


def robust_grade_manuscript(manuscript_text: str) -> Tuple[List[float], float, int, str, str]:
    """
    Grade manuscript, retrying if the scores are not in ascending order.
    It will attempt to grade the manuscript up to MAX_ATTEMPTS times.

    Returns:
        Tuple of (final_scores, final_weighted_score, num_attempts, decision_reason, justification)
    """
    print("    Starting robust grading process...")

    last_scores = [0] * 6
    last_justification = "Grading failed"

    for attempt in range(1, MAX_ATTEMPTS + 1):
        print(f"      Attempt {attempt}/{MAX_ATTEMPTS}...")
        scores, justification = grade_manuscript(manuscript_text)

        last_scores = scores  # Keep track of the last attempt
        last_justification = justification

        # Check 1: Was the API call successful and did it return scores?
        if not scores or all(s is None for s in scores):
            print("      ✗ API call failed or returned no scores. Retrying...")
            continue  # Move to the next attempt

        # Check 2: Are the scores in ascending order?
        if are_scores_ascending(scores):
            print("      ✓ Grading successful and scores are in ascending order.")
            weighted_score = calculate_weighted_score(scores)
            final_scores = [score if score is not None else 0 for score in scores]
            return final_scores, weighted_score, attempt, f"Successful on attempt {attempt}", justification
        else:
            print(f"      ✗ Scores are not in ascending order: {scores}. Retrying...")

    # If loop finishes, all attempts failed. Return the last result.
    print("      ✗ Max retries reached. Returning last failed attempt.")
    final_scores = [score if score is not None else 0 for score in last_scores]
    weighted_score = calculate_weighted_score(final_scores)
    decision_reason = "Failed: Scores not ascending after multiple retries"

    return final_scores, weighted_score, MAX_ATTEMPTS, decision_reason, last_justification


def process_single_manuscript(file_path):
    """Process a single manuscript file and return results"""
    print(f"\nProcessing: {os.path.basename(file_path)}")

    # Extract text from manuscript
    print(f"  Extracting text from manuscript...")
    manuscript_text = extract_text_from_file(file_path)
    if not manuscript_text:
        print(f"  ERROR: Failed to extract text from {file_path}")
        return None

    print(f"  Extracted {len(manuscript_text)} characters from manuscript")

    try:
        # Grade the manuscript
        final_scores, final_weighted_score, num_attempts, decision_reason, justification = robust_grade_manuscript(
            manuscript_text)

        print(f"  COMPLETED! Final Score: {final_weighted_score:.1f} (after {num_attempts} attempts)")
        print(f"    Decision: {decision_reason}")

        return {
            'filename': os.path.basename(file_path),
            'level1': final_scores[0],
            'level2': final_scores[1],
            'level3': final_scores[2],
            'level4': final_scores[3],
            'level5': final_scores[4],
            'level6': final_scores[5],
            'final_score': final_weighted_score,
            'num_attempts': num_attempts,
            'decision_reason': decision_reason,
            'justification': justification
        }

    except Exception as e:
        print(f"  ERROR: An unexpected error occurred: {e}")
        return None


def save_results_to_excel(results_list, output_file):
    """Save all results to Excel file"""
    try:
        # Create DataFrame
        df = pd.DataFrame(results_list)

        # Reorder columns
        column_order = ['filename', 'level1', 'level2', 'level3', 'level4', 'level5', 'level6',
                        'final_score', 'num_attempts', 'decision_reason', 'justification']
        df = df[column_order]

        # Round numerical columns to 1 decimal place
        numerical_cols = ['level1', 'level2', 'level3', 'level4', 'level5', 'level6', 'final_score']
        for col in numerical_cols:
            df[col] = df[col].round(1)

        # Save to Excel
        df.to_excel(output_file, index=False)
        print(f"\nResults saved to {output_file}")
        return True
    except Exception as e:
        print(f"Error saving results to Excel: {e}")
        return False





def main():
    """Main function to process all manuscripts in specified input folder."""
    args = parse_args()
    input_folder = args.input_folder
    output_file = args.output_file

    print("=== Batch Manuscript Grader ===")
    print(f"Input folder: {input_folder}")
    print(f"Output file: {output_file}")
    print(f"Grading model: {GRADING_MODEL}")
    print(f"Temperature: {TEMPERATURE}")
    print(f"Max Retries: {MAX_ATTEMPTS}")
    print("Processing: Sequential (1 second delay between API calls)")
    print("Supported file types: .docx, .txt, .pdf")

    # Check if input folder exists
    if not os.path.exists(input_folder):
        print(f"ERROR: Input folder '{input_folder}' does not exist!")
        print("Please create the folder and place your .docx, .txt, or .pdf files in it.")
        return

    # Find all supported manuscript files in input folder
    supported_extensions = {'.docx', '.txt', '.pdf'}
    manuscript_files = [
        os.path.join(input_folder, f)
        for f in os.listdir(input_folder)
        if os.path.splitext(f)[1].lower() in supported_extensions and not f.startswith('~')
    ]

    if not manuscript_files:
        print(f"ERROR: No .docx, .txt, or .pdf files found in '{input_folder}' folder!")
        return

    print(f"\nFound {len(manuscript_files)} manuscript files to process:")
    for file_path in manuscript_files:
        ext = os.path.splitext(file_path)[1].upper()
        print(f"  - {os.path.basename(file_path)} [{ext}]")

    # Process each file sequentially
    results_list = []
    start_time = datetime.now()
    print(f"\n{'=' * 60}")
    print("Starting sequential processing...")

    for i, file_path in enumerate(manuscript_files, start=1):
        print(f"\n[{i}/{len(manuscript_files)}] {os.path.basename(file_path)}")
        try:
            result = process_single_manuscript(file_path)
            if result:
                results_list.append(result)
                print(f"✓ Successfully completed {os.path.basename(file_path)}")
            else:
                print(f"✗ Failed to process {os.path.basename(file_path)}")
        except Exception as e:
            print(f"✗ Exception processing {os.path.basename(file_path)}: {e}")

    # Save and summarise results
    if results_list:
        print(f"\n{'=' * 60}")
        print(f"Batch processing completed: {len(results_list)} succeeded, {len(manuscript_files) - len(results_list)} failed")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_with_ts = Path(output_file)
        output_with_ts = output_with_ts.with_name(f"{output_with_ts.stem}_{timestamp}{output_with_ts.suffix}")
        if save_results_to_excel(results_list, str(output_with_ts)):
            print(f"\nResults saved to {output_with_ts}")
        duration = datetime.now() - start_time
        print(f"Total processing time: {duration}")
        avg = duration.total_seconds() / len(results_list)
        print(f"Average time per manuscript: {avg:.1f} seconds")
        total_calls = sum(r['num_attempts'] for r in results_list)
        print(f"Total API calls made: {total_calls}")
    else:
        print("ERROR: No files were successfully processed!")

if __name__ == "__main__":
    main()
